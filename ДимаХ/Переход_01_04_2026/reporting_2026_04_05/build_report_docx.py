#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Преобразование HTML-отчета (карточки) в DOCX для импорта в Confluence.

Альбомная ориентация, компактные шрифты, карточки в две колонки (как .grid в HTML),
широкие карточки grid-column 1 / -1 на всю строку. Таблицы, заливки, границы.
"""
from __future__ import annotations

import pathlib
import re
import sys

from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.table import _Cell
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BADGE_HEX = {
    "done": "28A745",
    "close": "F59E0B",
    "not-done": "DC3545",
    "pending": "1D4FA0",
    "no-data": "8A94A6",
}

DELTA_FILL = {
    "": "ECFEFF",
    "warn": "FFF7ED",
    "bad": "FFF1F2",
    "grey": "F8FAFC",
    "release": "FFF5F5",
}

# Tsveta granits kak v CSS (delta, kartochki, meta)
DELTA_BORDER = {
    "": "E2E8F0",
    "warn": "FED7AA",
    "bad": "FCA5A5",
    "grey": "E2E8F0",
    "release": "F5C0C0",
}

BORDER_DEFAULT = "E2E8F0"
BORDER_CARD_OUTER = "DBEAFE"
BORDER_SECTION_RESULTS = "BFDBFE"
BORDER_SECTION_ACTIONS = "BBF7D0"
BORDER_SOURCE = "DBEAFE"

# Shrift (kompaktno, pod albomnyy list ~ kak HTML 12--13px tekst)
FS = {
    "h1": Pt(14),
    "header_p": Pt(8.5),
    "source": Pt(8.5),
    "kpi_v": Pt(12),
    "kpi_l": Pt(7.5),
    "section": Pt(10),
    "plan_title": Pt(10),
    "base_label": Pt(8),
    "table_wrap_h2": Pt(11),
    "card_title": Pt(9),
    "badge": Pt(7),
    "t_lbl": Pt(6.5),
    "t_val": Pt(9.5),
    "t_sub": Pt(7.5),
    "delta": Pt(8),
    "meta": Pt(7),
    "meta_th_tgt": Pt(6.5),
    "meta_td_tgt": Pt(8.5),
    "main_th": Pt(7.5),
    "main_th_goal": Pt(6.5),
    "main_td": Pt(8),
    "main_td_goal": Pt(9.5),
    "main_td_goal_muted": Pt(7.5),
    "research_h2": Pt(11),
    "research_group": Pt(9),
    "research_icon": Pt(7),
    "research_body": Pt(8.5),
    "research_sub": Pt(8),
    "summary_h2": Pt(11),
    "summary_li": Pt(8.5),
}


def content_width_inches(section) -> float:
    """Shirina teksta mezhdu polyami."""
    return (
        section.page_width.inches
        - section.left_margin.inches
        - section.right_margin.inches
    )


def half_column_inches(section) -> float:
    """Polovina kontenta kak v HTML grid 2 kolonki (s zazorom)."""
    return max((content_width_inches(section) - 0.34) / 2, 2.85)


def full_card_inches(section) -> float:
    """Shirina kartochki na vsyu stroku (grid-column 1 / -1)."""
    return max(content_width_inches(section) - 0.04, 5.5)


def card_inner_width_inches(slot_inches: float) -> float:
    """Shirina teksta vnutri kartochki: vneshnie otstupy yacheyki + zapas."""
    return max(slot_inches - 0.14, 1.35)


def summary_table_inner_width(section) -> float:
    """Shirina svodnoy tablitsy v bloke table-wrap (s uchotom padding yacheyki)."""
    return max(
        full_card_inches(section) - 2 * (120 / 1440.0) - 0.1,
        5.0,
    )


# 7 kolonok svodnoy tablitsy (summa ~ 1)
MAIN_TABLE_COL_WEIGHTS = [0.056, 0.188, 0.128, 0.128, 0.158, 0.118, 0.124]


def block_is_full_width(block) -> bool:
    st = block.get("style", "") or ""
    return "grid-column" in st and "1 / -1" in st.replace(" ", "")


def _border_ooxml(side: str, val: str, sz_eighths: int, color: str) -> OxmlElement:
    el = OxmlElement(f"w:{side}")
    el.set(qn("w:val"), val)
    el.set(qn("w:sz"), str(sz_eighths))
    el.set(qn("w:space"), "0")
    el.set(qn("w:color"), color)
    return el


def set_table_borders(
    table,
    color: str = BORDER_DEFAULT,
    sz_eighths: int = 8,
) -> None:
    """Edinaya setka: vne i vnutri (Word: sz v 1/8 punkta, 8 ~ 1 pt)."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    old = tbl_pr.find(qn("w:tblBorders"))
    if old is not None:
        tbl_pr.remove(old)
    tb = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tb.append(_border_ooxml(name, "single", sz_eighths, color))
    tbl_pr.append(tb)


def set_cell_borders(
    cell,
    top=None,
    bottom=None,
    left=None,
    right=None,
) -> None:
    """Kazhdyy parametp: None ili (sz_eighths, color_hex)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    old = tc_pr.find(qn("w:tcBorders"))
    if old is not None:
        tc_pr.remove(old)
    tb = OxmlElement("w:tcBorders")
    mapping = (
        ("top", top),
        ("bottom", bottom),
        ("left", left),
        ("right", right),
    )
    for side, spec in mapping:
        if spec is None:
            continue
        sz, col = spec
        tb.append(_border_ooxml(side, "single", sz, col))
    if len(tb):
        tc_pr.append(tb)


def _border_none(side: str) -> OxmlElement:
    el = OxmlElement(f"w:{side}")
    el.set(qn("w:val"), "none")
    el.set(qn("w:sz"), "0")
    el.set(qn("w:space"), "0")
    el.set(qn("w:color"), "auto")
    return el


def set_table_no_outer_border(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    old = tbl_pr.find(qn("w:tblBorders"))
    if old is not None:
        tbl_pr.remove(old)
    tb = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tb.append(_border_none(name))
    tbl_pr.append(tb)


def set_cell_bottom_separator(cell, sz_eighths: int = 8, color: str = BORDER_DEFAULT) -> None:
    """Tolko nizhnyaya liniya (shapka kartochki kak v HTML .top)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    old = tc_pr.find(qn("w:tcBorders"))
    if old is not None:
        tc_pr.remove(old)
    tb = OxmlElement("w:tcBorders")
    for s in ("top", "left", "right"):
        tb.append(_border_none(s))
    tb.append(_border_ooxml("bottom", "single", sz_eighths, color))
    tc_pr.append(tb)


def set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120) -> None:
    tc = cell._tc
    tc_mar = OxmlElement("w:tcMar")
    for side, val in (
        ("top", top),
        ("bottom", bottom),
        ("left", left),
        ("right", right),
    ):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tc_mar.append(el)
    tc_pr = tc.get_or_add_tcPr()
    tc_pr.append(tc_mar)


def set_table_width_pct(table, pct_hundred: int = 100) -> None:
    """Shirina tablitsy v %% ot konteynera (5000 = 100%% v OOXML pct)."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    old = tbl_pr.find(qn("w:tblW"))
    if old is not None:
        tbl_pr.remove(old)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(pct_hundred * 50))
    tbl_w.set(qn("w:type"), "pct")
    tbl_pr.append(tbl_w)


def parse_rgb_from_style(style: str) -> RGBColor | None:
    if not style:
        return None
    m = re.search(r"color:\s*#([0-9a-fA-F]{6})", style)
    if m:
        h = m.group(1)
        return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))
    m2 = re.search(
        r"color:\s*rgb\s*\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*\)", style
    )
    if m2:
        return RGBColor(int(m2.group(1)), int(m2.group(2)), int(m2.group(3)))
    return None


def paragraph_clear(paragraph) -> None:
    paragraph.text = ""


def add_inline_runs(paragraph, element) -> None:
    if element is None:
        return

    def walk(node):
        if isinstance(node, NavigableString):
            paragraph.add_run(str(node).replace("\xa0", " "))
            return
        if not hasattr(node, "name") or node.name is None:
            return
        if node.name == "br":
            paragraph.add_run().add_break()
            return
        if node.name == "b":
            r = paragraph.add_run(node.get_text())
            r.bold = True
            return
        if node.name == "span":
            r = paragraph.add_run(node.get_text())
            rgb = parse_rgb_from_style(node.get("style", ""))
            if rgb:
                r.font.color.rgb = rgb
            st = node.get("style", "")
            r.bold = "font-weight:700" in st or "font-weight:bold" in st
            return
        for ch in node.children:
            walk(ch)

    walk(element)


def badge_class_from_element(badge_el) -> str:
    classes = badge_el.get("class", [])
    for c in BADGE_HEX:
        if c in classes:
            return c
    return "no-data"


def add_card_at(
    parent,
    block,
    width_inches: float,
    trailing_para: bool = False,
) -> None:
    slot = max(width_inches - 0.03, 1.35)
    body_w = card_inner_width_inches(slot)
    outer = parent.add_table(rows=1, cols=1)
    outer.autofit = False
    outer.alignment = WD_TABLE_ALIGNMENT.CENTER
    outer.columns[0].width = Inches(slot)
    set_table_borders(outer, BORDER_CARD_OUTER, 6)
    set_table_width_pct(outer, 100)
    oc = outer.rows[0].cells[0]
    set_cell_margins(oc, top=24, bottom=24, left=32, right=32)
    set_cell_shading(oc, "FFFFFF")
    inner_doc = oc

    top_el = block.select_one(".top")
    title_el = top_el.select_one(".title") if top_el else None
    badge_el = top_el.select_one(".badge") if top_el else None

    hdr = inner_doc.add_table(rows=1, cols=2)
    hdr.autofit = True
    set_table_no_outer_border(hdr)
    c0, c1 = hdr.rows[0].cells
    c0.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    c1.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(c0, "FFFFFF")
    set_cell_bottom_separator(c0, 8, BORDER_DEFAULT)
    set_cell_bottom_separator(c1, 8, BORDER_DEFAULT)
    p0 = c0.paragraphs[0]
    p0.paragraph_format.space_after = Pt(0)
    if title_el:
        r = p0.add_run(title_el.get_text(strip=True))
        r.bold = True
        r.font.size = FS["card_title"]
        r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p1.paragraph_format.space_after = Pt(0)
    if badge_el:
        btext = badge_el.get_text(strip=True)
        bcls = badge_class_from_element(badge_el)
        run = p1.add_run(" " + btext + " ")
        run.bold = True
        run.font.size = FS["badge"]
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(c1, BADGE_HEX[bcls])
    for brdr in hdr.rows[0].cells:
        set_cell_margins(brdr, top=22, bottom=22, left=28, right=28)

    content_el = block.select_one(".content")
    if not content_el:
        if trailing_para and not isinstance(parent, _Cell):
            parent.add_paragraph()
        return

    times_el = content_el.select_one(".times")
    if times_el:
        tboxes = times_el.select(".tbox")
        tt = inner_doc.add_table(rows=1, cols=max(1, len(tboxes)))
        tt.autofit = False
        set_table_borders(tt, BORDER_DEFAULT, 6)
        set_table_width_pct(tt, 100)
        ntb = len(tboxes)
        if ntb == 2:
            tt.columns[0].width = Inches(body_w * 0.49)
            tt.columns[1].width = Inches(body_w * 0.49)
        elif ntb == 1:
            tt.columns[0].width = Inches(body_w * 0.98)
        for i, tbox in enumerate(tboxes):
            cell = tt.rows[0].cells[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_shading(cell, "F8FAFC")
            set_cell_margins(cell, top=28, bottom=28, left=28, right=28)
            st_box = tbox.get("style", "")
            if "fff1f2" in st_box.replace(" ", "").lower():
                set_cell_shading(cell, "FFF1F2")
            if "fff7ed" in st_box.replace(" ", "").lower():
                set_cell_shading(cell, "FFF7ED")
            lbl = tbox.select_one(".lbl")
            val = tbox.select_one(".val")
            sub = tbox.select_one(".sub")
            cl = tbox.get("class", [])
            is_old = "old" in cl
            is_new = "new" in cl
            if lbl:
                pl = cell.add_paragraph()
                pl.paragraph_format.space_after = Pt(2)
                rl = pl.add_run(lbl.get_text(strip=True))
                rl.font.size = FS["t_lbl"]
                rl.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
                rl.bold = True
            if val:
                pv = cell.add_paragraph()
                pv.paragraph_format.space_after = Pt(2)
                add_inline_runs(pv, val)
                for r in pv.runs:
                    r.font.size = FS["t_val"]
                    r.bold = True
                if is_old:
                    for r in pv.runs:
                        r.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
                elif is_new:
                    for r in pv.runs:
                        r.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)
            if sub:
                ps = cell.add_paragraph()
                add_inline_runs(ps, sub)
                for r in ps.runs:
                    r.font.size = FS["t_sub"]
                    r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    deltas = content_el.find_all(
        "div",
        class_=lambda c: bool(c) and "delta" in c,
        recursive=False,
    )
    for delta_el in deltas:
        dclasses = delta_el.get("class", [])
        dkey = ""
        if "warn" in dclasses:
            dkey = "warn"
        elif "bad" in dclasses:
            dkey = "bad"
        elif "grey" in dclasses:
            dkey = "grey"
        elif "release" in dclasses:
            dkey = "release"
        dt = inner_doc.add_table(rows=1, cols=1)
        dc = dt.rows[0].cells[0]
        brd = DELTA_BORDER.get(dkey, DELTA_BORDER[""])
        set_table_borders(dt, brd, 6)
        set_table_width_pct(dt, 100)
        set_cell_shading(dc, DELTA_FILL.get(dkey, DELTA_FILL[""]))
        set_cell_margins(dc, top=32, bottom=32, left=32, right=32)
        if dkey == "release":
            set_cell_borders(
                dc,
                top=(8, brd),
                bottom=(8, brd),
                right=(8, brd),
                left=(28, "DC3545"),
            )
        dp = dc.paragraphs[0]
        dp.paragraph_format.space_after = Pt(0)
        paragraph_clear(dp)
        add_inline_runs(dp, delta_el)
        for r in dp.runs:
            r.font.size = FS["delta"]
            if dkey == "release":
                r.font.color.rgb = RGBColor(0x7A, 0x1A, 0x1A)
                r.bold = True

    meta_tables = content_el.find_all(
        "table",
        class_=lambda c: bool(c) and "meta-table" in c,
        recursive=False,
    )
    for mt in meta_tables:
        rows = mt.select("tr")
        if not rows:
            continue
        mtab = inner_doc.add_table(rows=len(rows), cols=2)
        mtab.autofit = False
        set_table_borders(mtab, BORDER_DEFAULT, 6)
        set_table_width_pct(mtab, 100)
        mtab.columns[0].width = Inches(body_w * 0.32)
        mtab.columns[1].width = Inches(body_w * 0.66)
        for ri, tr in enumerate(rows):
            th = tr.find("th")
            td = tr.find("td")
            c_a = mtab.rows[ri].cells[0]
            c_b = mtab.rows[ri].cells[1]
            tr_class = tr.get("class", [])
            if "meta-target" in tr_class:
                set_cell_shading(c_a, "E0E7FF")
                set_cell_shading(c_b, "F5F3FF")
                set_cell_borders(
                    c_a,
                    top=(18, "C7D2FE"),
                    left=(28, "4338CA"),
                    bottom=(8, BORDER_DEFAULT),
                    right=(8, BORDER_DEFAULT),
                )
                set_cell_borders(
                    c_b,
                    top=(18, "C7D2FE"),
                    bottom=(8, BORDER_DEFAULT),
                    left=(8, BORDER_DEFAULT),
                    right=(8, BORDER_DEFAULT),
                )
            else:
                set_cell_shading(c_a, "F8FAFC")
                set_cell_shading(c_b, "FFFFFF")
            paragraph_clear(c_a.paragraphs[0])
            paragraph_clear(c_b.paragraphs[0])
            if th:
                add_inline_runs(c_a.paragraphs[0], th)
                if "meta-target" in tr_class:
                    for r in c_a.paragraphs[0].runs:
                        r.font.color.rgb = RGBColor(0x31, 0x2E, 0x81)
                        r.bold = True
            if td:
                add_inline_runs(c_b.paragraphs[0], td)
                if "meta-target" in tr_class:
                    for r in c_b.paragraphs[0].runs:
                        r.font.color.rgb = RGBColor(0x1E, 0x1B, 0x4B)
                        r.bold = True
                    if td.get("class") and "muted-target" in td.get("class", []):
                        for r in c_b.paragraphs[0].runs:
                            r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
                            r.bold = True
                            r.italic = True
            if "meta-target" not in tr_class:
                for r in c_a.paragraphs[0].runs:
                    r.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
                    r.font.size = FS["meta"]
                for r in c_b.paragraphs[0].runs:
                    r.font.size = FS["meta"]
            else:
                for r in c_a.paragraphs[0].runs:
                    r.font.size = FS["meta_th_tgt"]
                for r in c_b.paragraphs[0].runs:
                    r.font.size = FS["meta_td_tgt"]

    if trailing_para and not isinstance(parent, _Cell):
        parent.add_paragraph()


def _flush_grid_pair(doc: Document, left_block, right_block, section) -> None:
    pair = doc.add_table(1, 2)
    pair.autofit = False
    w = half_column_inches(section)
    pair.columns[0].width = Inches(w)
    pair.columns[1].width = Inches(w)
    set_table_no_outer_border(pair)
    cl, cr = pair.rows[0].cells
    cl.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    cr.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_margins(cl, right=18)
    set_cell_margins(cr, left=18)
    add_card_at(cl, left_block, w, trailing_para=False)
    add_card_at(cr, right_block, w, trailing_para=False)
    doc.add_paragraph()


def _flush_grid_single(doc: Document, block, section) -> None:
    pair = doc.add_table(1, 2)
    pair.autofit = False
    w = half_column_inches(section)
    pair.columns[0].width = Inches(w)
    pair.columns[1].width = Inches(w)
    set_table_no_outer_border(pair)
    cl, cr = pair.rows[0].cells
    cl.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    cr.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    add_card_at(cl, block, w, trailing_para=False)
    paragraph_clear(cr.paragraphs[0])
    doc.add_paragraph()


def process_grid(doc: Document, grid_el, section) -> None:
    blocks = list(
        grid_el.find_all(
            "div",
            class_=lambda c: bool(c) and "block" in c,
            recursive=False,
        )
    )
    i = 0
    while i < len(blocks):
        b = blocks[i]
        if block_is_full_width(b):
            add_card_at(doc, b, full_card_inches(section), trailing_para=True)
            i += 1
            continue
        if i + 1 < len(blocks) and not block_is_full_width(blocks[i + 1]):
            _flush_grid_pair(doc, blocks[i], blocks[i + 1], section)
            i += 2
        else:
            _flush_grid_single(doc, b, section)
            i += 1


def add_main_table(parent, table_el, section=None) -> None:
    rows = table_el.select("tr")
    if not rows:
        return
    ncol = len(rows[0].find_all(["th", "td"]))
    tbl = parent.add_table(rows=len(rows), cols=ncol)
    tbl.autofit = False
    set_table_borders(tbl, BORDER_DEFAULT, 6)
    set_table_width_pct(tbl, 100)
    if section is not None and ncol == len(MAIN_TABLE_COL_WEIGHTS):
        inner_w = summary_table_inner_width(section)
        for ci, wt in enumerate(MAIN_TABLE_COL_WEIGHTS):
            tbl.columns[ci].width = Inches(inner_w * wt)
    for ri, tr in enumerate(rows):
        cells = tr.find_all(["th", "td"])
        for ci, cell_el in enumerate(cells):
            c = tbl.rows[ri].cells[ci]
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            paragraph_clear(c.paragraphs[0])
            add_inline_runs(c.paragraphs[0], cell_el)
            cls = cell_el.get("class", [])

            if cell_el.name == "th":
                set_cell_shading(c, "EEF2FF" if ci == 4 else "F8FAFC")
                if ci == 4:
                    for r in c.paragraphs[0].runs:
                        r.bold = True
                        r.font.color.rgb = RGBColor(0x31, 0x2E, 0x81)
                        r.font.size = FS["main_th_goal"]
                    set_cell_borders(
                        c,
                        bottom=(20, "818CF8"),
                        top=(8, BORDER_DEFAULT),
                        left=(8, BORDER_DEFAULT),
                        right=(8, BORDER_DEFAULT),
                    )
                else:
                    for r in c.paragraphs[0].runs:
                        r.bold = True
                        r.font.size = FS["main_th"]

            elif cell_el.name == "td":
                if ci == 4:
                    if "muted" in cls:
                        set_cell_shading(c, "F1F5F9")
                        set_cell_borders(
                            c,
                            left=(18, "CBD5E1"),
                            top=(8, BORDER_DEFAULT),
                            bottom=(8, BORDER_DEFAULT),
                            right=(8, BORDER_DEFAULT),
                        )
                        for r in c.paragraphs[0].runs:
                            r.bold = True
                            r.font.size = FS["main_td_goal_muted"]
                            r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
                    else:
                        set_cell_shading(c, "FAF5FF")
                        set_cell_borders(
                            c,
                            left=(24, "7C3AED"),
                            top=(8, BORDER_DEFAULT),
                            bottom=(8, BORDER_DEFAULT),
                            right=(8, BORDER_DEFAULT),
                        )
                        for r in c.paragraphs[0].runs:
                            r.bold = True
                            r.font.size = FS["main_td_goal"]
                            r.font.color.rgb = RGBColor(0x37, 0x30, 0xA3)
                else:
                    for r in c.paragraphs[0].runs:
                        r.font.size = FS["main_td"]
                    if "ok" in cls:
                        for r in c.paragraphs[0].runs:
                            r.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)
                            r.bold = True
                    elif "bad" in cls:
                        for r in c.paragraphs[0].runs:
                            r.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
                            r.bold = True
                    elif "warn" in cls:
                        for r in c.paragraphs[0].runs:
                            r.font.color.rgb = RGBColor(0xB4, 0x53, 0x09)
                            r.bold = True
                    elif "muted" in cls:
                        for r in c.paragraphs[0].runs:
                            r.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    if not isinstance(parent, _Cell):
        parent.add_paragraph()


def add_research_card(doc: Document, card) -> None:
    classes = card.get("class", [])
    fill = "FFFBF0"
    b_line = "F0C940"
    if "in-progress" in classes:
        fill = "F0F9FF"
        b_line = "7DD3FC"
    elif "planned" in classes:
        fill = "F8FAFC"
        b_line = "CBD5E1"

    t = doc.add_table(rows=1, cols=2)
    t.autofit = False
    set_table_borders(t, b_line, 8)
    set_table_width_pct(t, 100)
    cw = content_width_inches(doc.sections[0])
    icon_w = 0.82
    body_w = max(cw - icon_w - 0.14, 3.2)
    t.columns[0].width = Inches(icon_w)
    t.columns[1].width = Inches(body_w)
    ic, bc = t.rows[0].cells
    icon_el = card.select_one(".rc-icon")
    if icon_el:
        if "active" in icon_el.get("class", []):
            set_cell_shading(ic, "0EA5E9")
        elif "idea" in icon_el.get("class", []):
            set_cell_shading(ic, "E0A800")
        else:
            set_cell_shading(ic, "94A3B8")
        paragraph_clear(ic.paragraphs[0])
        ip = ic.paragraphs[0]
        ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ir = ip.add_run(icon_el.get_text(strip=True))
        ir.font.size = FS["research_icon"]
        ir.bold = True
        ir.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_shading(bc, fill)
    set_cell_margins(bc, top=36, bottom=36, left=40, right=40)
    set_cell_margins(ic, top=36, bottom=36, left=28, right=28)

    body = card.select_one(".rc-body")
    if body:
        first = True
        for sub in body.children:
            if isinstance(sub, NavigableString):
                if str(sub).strip():
                    bp = bc.paragraphs[0] if first else bc.add_paragraph()
                    first = False
                    trn = bp.add_run(str(sub).strip())
                    trn.font.size = FS["research_body"]
                continue
            if sub.name == "div" and "rc-subtask" in sub.get("class", []):
                sp = bc.add_paragraph()
                first = False
                add_inline_runs(sp, sub)
                for r in sp.runs:
                    r.font.size = FS["research_sub"]
                    r.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
                continue
            if sub.name:
                bp = bc.paragraphs[0] if first else bc.add_paragraph()
                first = False
                paragraph_clear(bp)
                add_inline_runs(bp, sub)
                for r in bp.runs:
                    r.font.size = FS["research_body"]
    doc.add_paragraph()


def add_research_section(doc: Document, section_el) -> None:
    h2 = section_el.find("h2")
    if h2:
        p = doc.add_paragraph()
        r = p.add_run(h2.get_text(strip=True))
        r.bold = True
        r.font.size = FS["research_h2"]
    for rg in section_el.select(".research-group"):
        gt = rg.select_one(".research-group-title")
        if gt:
            gp = doc.add_paragraph()
            gr = gp.add_run(gt.get_text(strip=True))
            gr.bold = True
            gr.font.size = FS["research_group"]
            gr.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        for card in rg.select(".research-card"):
            add_research_card(doc, card)


def convert(html_path: pathlib.Path, docx_path: pathlib.Path) -> None:
    raw = html_path.read_text(encoding="utf-8")
    soup = BeautifulSoup(raw, "lxml")
    doc = Document()

    sec = doc.sections[0]
    sec.left_margin = Inches(0.45)
    sec.right_margin = Inches(0.45)
    sec.top_margin = Inches(0.45)
    sec.bottom_margin = Inches(0.45)
    new_width, new_height = sec.page_height, sec.page_width
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = new_width
    sec.page_height = new_height

    wrap = soup.select_one(".wrap")
    if not wrap:
        raise SystemExit("No .wrap in HTML")

    header_el = wrap.select_one(".header")
    if header_el:
        ht = doc.add_table(rows=1, cols=1)
        ht.autofit = False
        ht.columns[0].width = Inches(full_card_inches(sec))
        hc = ht.rows[0].cells[0]
        set_cell_shading(hc, "1D4FA0")
        set_cell_margins(hc, top=120, bottom=120, left=140, right=140)
        h1 = hc.paragraphs[0]
        h1t = header_el.find("h1")
        if h1t:
            r = h1.add_run(h1t.get_text(strip=True))
            r.bold = True
            r.font.size = FS["h1"]
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for pe in header_el.find_all("p"):
            hp = hc.add_paragraph()
            add_inline_runs(hp, pe)
            for r in hp.runs:
                r.font.size = FS["header_p"]
                r.font.color.rgb = RGBColor(0xEE, 0xF2, 0xFF)
        doc.add_paragraph()

    src_el = wrap.select_one(".source-box")
    if src_el:
        st = doc.add_table(rows=1, cols=1)
        set_table_borders(st, BORDER_SOURCE, 8)
        sc = st.rows[0].cells[0]
        set_cell_shading(sc, "FFFFFF")
        set_cell_margins(sc, top=80, bottom=80, left=100, right=100)
        sp = sc.paragraphs[0]
        paragraph_clear(sp)
        add_inline_runs(sp, src_el)
        for r in sp.runs:
            r.font.size = FS["source"]
        doc.add_paragraph()

    kpi_el = wrap.select_one(".kpi")
    if kpi_el:
        items = kpi_el.select(".kpi-item")
        kt = doc.add_table(rows=1, cols=len(items))
        kt.autofit = False
        kw = max((content_width_inches(sec) - 0.45) / len(items), 0.98)
        for idx in range(len(items)):
            kt.columns[idx].width = Inches(kw)
        set_table_borders(kt, BORDER_DEFAULT, 8)
        for i, it in enumerate(items):
            cell = kt.rows[0].cells[i]
            set_cell_shading(cell, "FFFFFF")
            set_cell_margins(cell, top=44, bottom=44, left=28, right=28)
            v = it.select_one(".v")
            le = it.select_one(".l")
            if v:
                vp = cell.add_paragraph()
                vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_inline_runs(vp, v)
                for r in vp.runs:
                    r.bold = True
                    r.font.size = FS["kpi_v"]
            if le:
                lp = cell.add_paragraph()
                lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                lr = lp.add_run(le.get_text(strip=True))
                lr.font.size = FS["kpi_l"]
                lr.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        doc.add_paragraph()

    zone_res = soup.select_one(".zone-results")
    if zone_res:
        for child in zone_res.children:
            if not getattr(child, "name", None):
                continue
            cl = child.get("class", [])
            if "section-title" in cl and "results" in cl:
                tt = doc.add_table(rows=1, cols=1)
                set_table_borders(tt, BORDER_SECTION_RESULTS, 8)
                tc = tt.rows[0].cells[0]
                set_cell_shading(tc, "EFF6FF")
                set_cell_margins(tc, top=70, bottom=70, left=100, right=100)
                main_t = child.get_text(" ", strip=True)
                main_t = re.sub(r"\s+", " ", main_t)
                p = tc.paragraphs[0]
                r = p.add_run(main_t)
                r.bold = True
                r.font.size = FS["section"]
                r.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
                doc.add_paragraph()
            elif "base-label" in cl:
                bt = doc.add_table(rows=1, cols=1)
                set_table_borders(bt, "0E7A5F", 10)
                blc = bt.rows[0].cells[0]
                set_cell_margins(blc, top=35, bottom=35, left=80, right=80)
                set_cell_shading(blc, "FFFFFF")
                bp = blc.paragraphs[0]
                br = bp.add_run(child.get_text(" ", strip=True))
                br.bold = True
                br.font.size = FS["base_label"]
                br.font.color.rgb = RGBColor(0x0E, 0x7A, 0x5F)
                doc.add_paragraph()
            elif "grid" in cl:
                process_grid(doc, child, sec)
            elif "table-wrap" in cl:
                h2 = child.find("h2")
                tbl = child.find("table")
                wrap_tbl = doc.add_table(rows=1, cols=1)
                wrap_tbl.autofit = False
                wrap_tbl.columns[0].width = Inches(full_card_inches(sec))
                set_table_borders(wrap_tbl, BORDER_DEFAULT, 8)
                wcell = wrap_tbl.rows[0].cells[0]
                set_cell_shading(wcell, "FFFFFF")
                set_cell_margins(wcell, top=72, bottom=72, left=72, right=72)
                if h2:
                    hp = wcell.add_paragraph()
                    hr = hp.add_run(h2.get_text(strip=True))
                    hr.bold = True
                    hr.font.size = FS["table_wrap_h2"]
                if tbl:
                    add_main_table(wcell, tbl, sec)
                doc.add_paragraph()

    zone_act = soup.select_one(".zone-actions")
    if zone_act:
        det = zone_act.find("details")
        if det:
            summ = det.find("summary")
            if summ:
                tt = doc.add_table(rows=1, cols=1)
                set_table_borders(tt, BORDER_SECTION_ACTIONS, 8)
                tc = tt.rows[0].cells[0]
                set_cell_shading(tc, "F0FDF4")
                set_cell_margins(tc, top=70, bottom=70, left=100, right=100)
                tx = summ.get_text(" ", strip=True)
                tx = re.sub(r"\s+", " ", tx)
                p = tc.paragraphs[0]
                r = p.add_run(tx)
                r.bold = True
                r.font.size = FS["plan_title"]
                r.font.color.rgb = RGBColor(0x16, 0x65, 0x34)
                doc.add_paragraph()
            rs = det.select_one(".research-section")
            if rs:
                add_research_section(doc, rs)

    summary_el = wrap.select_one(".summary")
    if summary_el:
        h2 = summary_el.find("h2")
        if h2:
            sp = doc.add_paragraph()
            sr = sp.add_run(h2.get_text(strip=True))
            sr.bold = True
            sr.font.size = FS["summary_h2"]
        box = doc.add_table(rows=1, cols=1)
        set_table_borders(box, BORDER_DEFAULT, 8)
        bc = box.rows[0].cells[0]
        set_cell_shading(bc, "FFFFFF")
        set_cell_margins(bc, top=80, bottom=80, left=100, right=100)
        ul = summary_el.find("ul")
        if ul:
            for li in ul.find_all("li", recursive=False):
                lp = bc.add_paragraph()
                lp.paragraph_format.left_indent = Inches(0.2)
                lp.style = "List Bullet"
                paragraph_clear(lp)
                add_inline_runs(lp, li)
                for r in lp.runs:
                    r.font.size = FS["summary_li"]

    doc.save(str(docx_path))


def main() -> None:
    base = pathlib.Path(__file__).resolve().parent
    html = base / "report_cards_2026_04_05.html"
    out = base / "report_cards_2026_04_05.docx"
    if len(sys.argv) >= 2:
        html = pathlib.Path(sys.argv[1])
    if len(sys.argv) >= 3:
        out = pathlib.Path(sys.argv[2])
    convert(html, out)
    print("OK: " + str(out))


if __name__ == "__main__":
    main()
